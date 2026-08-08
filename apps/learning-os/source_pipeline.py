from __future__ import annotations

import csv
import hashlib
import json
import mimetypes
import re
import shutil
from datetime import UTC, datetime
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from ebooklib import ITEM_DOCUMENT, epub
from markdownify import markdownify as html_to_markdown
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".epub": "epub",
    ".docx": "docx",
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".html": "html",
    ".htm": "html",
    ".json": "json",
    ".csv": "csv",
    ".tsv": "tsv",
    ".url": "url-shortcut",
}

SUPPORTED_FORMAT_LABELS = (
    "PDF",
    "EPUB",
    "DOCX",
    "TXT",
    "Markdown",
    "HTML",
    "JSON",
    "CSV",
    "TSV",
    "Web URLs",
)

TEXT_PREVIEW_LIMIT = 6000
URL_USER_AGENT = "Lexicon/0.1"
HTML_MIME_PREFIXES = ("text/html", "application/xhtml+xml")
JSON_MIME_PREFIXES = ("application/json", "text/json")
TEXT_MIME_PREFIXES = ("text/plain", "text/markdown", "text/csv", "text/tab-separated-values")
PDF_MIME_PREFIXES = ("application/pdf",)


def iso_now() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def slugify(text: str, fallback: str = "source") -> str:
    cleaned = re.sub(r"[^a-z0-9]+", "-", str(text or "").lower()).strip("-")
    return cleaned[:60] or fallback


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_text(text: str) -> str:
    return sha256_bytes(text.encode("utf-8"))


def ensure_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def first_non_empty(*values: str) -> str:
    for value in values:
        candidate = str(value or "").strip()
        if candidate:
            return candidate
    return ""


def guess_format_from_path(path: Path) -> str:
    return SUPPORTED_EXTENSIONS.get(path.suffix.lower(), "")


def decode_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return ensure_text(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
    return ensure_text(raw.decode("utf-8", errors="ignore"))


def normalize_url(value: str) -> str:
    candidate = str(value or "").strip()
    if not candidate:
        raise ValueError("No URL provided")
    if not re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*://", candidate):
        candidate = f"https://{candidate}"
    parsed = urlparse(candidate)
    if not parsed.scheme or not parsed.netloc:
        raise ValueError("URL is not valid")
    return candidate


def html_to_text_preview(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for node in soup(["script", "style", "noscript"]):
        node.decompose()
    text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
    return text[:TEXT_PREVIEW_LIMIT]


def clean_markdown(markdown: str) -> str:
    text = markdown.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def markdown_table(rows: list[list[str]], max_rows: int = 10) -> str:
    if not rows:
        return "_No tabular rows found._"
    head = rows[0]
    sample = rows[1:max_rows]
    width = len(head)
    normalized_rows = [head] + [row + [""] * (width - len(row)) for row in sample]
    header = "| " + " | ".join(cell.strip() for cell in normalized_rows[0]) + " |"
    divider = "| " + " | ".join("---" for _ in normalized_rows[0]) + " |"
    body = [
        "| " + " | ".join(cell.strip() for cell in row[:width]) + " |"
        for row in normalized_rows[1:]
    ]
    return "\n".join([header, divider, *body]) if body else "\n".join([header, divider])


class SourcePipeline:
    def __init__(self, workspace_root: Path):
        self.workspace_root = Path(workspace_root)
        self.raw_root = self.workspace_root / "raw"
        self.processed_root = self.workspace_root / "processed"
        self.wiki_root = self.workspace_root / "wiki"
        for folder in (self.raw_root, self.processed_root, self.wiki_root):
            folder.mkdir(parents=True, exist_ok=True)

    def overview(self) -> dict:
        sources = self.list_sources()
        processed = [source for source in sources if source["status"] == "processed"]
        return {
            "app_name": "Lexicon.ai",
            "raw_count": len(sources),
            "processed_count": len(processed),
            "wiki_count": sum(1 for _ in self.wiki_root.rglob("*.md")),
            "supported_formats": list(SUPPORTED_FORMAT_LABELS),
            "sources": sources,
        }

    def list_sources(self) -> list[dict]:
        items: list[dict] = []
        for source_dir in sorted(self.raw_root.iterdir(), reverse=True) if self.raw_root.exists() else []:
            if not source_dir.is_dir():
                continue
            metadata = self._read_metadata(source_dir)
            if not metadata:
                continue
            items.append(self._summarize_metadata(metadata))
        items.sort(key=lambda item: item.get("imported_at", ""), reverse=True)
        return items

    def get_source(self, source_id: str) -> dict:
        metadata = self._read_metadata(self.raw_root / source_id)
        if not metadata:
            return {"ok": False, "error": f"Unknown source: {source_id}"}
        summary = self._summarize_metadata(metadata)
        processed_document = self._processed_document_path(source_id)
        processed_text = processed_document.read_text(encoding="utf-8") if processed_document.exists() else ""
        return {
            "ok": True,
            "source": {
                **summary,
                "metadata": metadata,
                "raw_preview": self._raw_preview(metadata),
                "processed_markdown": processed_text,
            },
        }

    def delete_source(self, source_id: str) -> dict:
        """Remove a source's raw artifacts and processed document. Deterministic
        and total — this is the raw-layer half of a cascade delete; the
        knowledge-build layer (manifest, note page, entities/relations) is the
        caller's responsibility via ``JobManager.forget_source``."""
        raw_dir = self.raw_root / source_id
        processed_dir = self.processed_root / source_id
        existed = raw_dir.exists() or processed_dir.exists()
        if raw_dir.exists():
            shutil.rmtree(raw_dir)
        if processed_dir.exists():
            shutil.rmtree(processed_dir)
        if not existed:
            return {"ok": False, "error": f"Unknown source: {source_id}"}
        return {"ok": True, "source_id": source_id}

    def import_files(self, paths: list[str]) -> dict:
        imported = []
        errors = []
        for raw_path in paths:
            try:
                source = self.import_file(Path(raw_path))
                imported.append(source)
            except Exception as exc:
                errors.append({"path": str(raw_path), "error": str(exc)})
        return {"ok": not errors, "sources": imported, "errors": errors}

    def import_file(self, path: Path) -> dict:
        path = Path(path)
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"File not found: {path}")
        source_format = guess_format_from_path(path)
        if source_format == "url-shortcut":
            url = self._parse_windows_url_shortcut(path)
            result = self.import_url(url)
            result["source"]["shortcut_name"] = path.name
            return result
        if not source_format:
            raise ValueError(f"Unsupported file format: {path.suffix or path.name}")

        raw_bytes = path.read_bytes()
        source_id = f"src-{sha256_bytes(raw_bytes)[:12]}"
        source_dir = self.raw_root / source_id
        existing = self._read_metadata(source_dir)
        if existing:
            return self.get_source(source_id)
        source_dir.mkdir(parents=True, exist_ok=True)
        original_name = path.name
        preserved_name = f"original{path.suffix.lower()}"
        raw_copy = source_dir / preserved_name
        shutil.copy2(path, raw_copy)

        metadata = {
            "id": source_id,
            "kind": "file",
            "format": source_format,
            "title": path.stem.replace("_", " ").strip() or path.stem,
            "original_name": original_name,
            "source_url": "",
            "canonical_url": "",
            "imported_at": iso_now(),
            "processed_at": "",
            "status": "imported",
            "sha256": sha256_bytes(raw_bytes),
            "mime_type": mimetypes.guess_type(original_name)[0] or "application/octet-stream",
            "raw_items": [{
                "label": "Original file",
                "path": str(raw_copy.relative_to(self.workspace_root)),
                "size": raw_copy.stat().st_size,
            }],
            "processing_error": "",
        }
        self._write_metadata(source_dir, metadata)
        return self.process_source(source_id)

    def import_url(self, value: str) -> dict:
        source_url = normalize_url(value)
        response = requests.get(
            source_url,
            timeout=30,
            headers={"User-Agent": URL_USER_AGENT},
        )
        response.raise_for_status()
        final_url = response.url or source_url
        content_type = (response.headers.get("content-type") or "").split(";")[0].strip().lower()
        source_key = final_url.encode("utf-8") + b"\0" + response.content
        source_id = f"src-{sha256_bytes(source_key)[:12]}"
        source_dir = self.raw_root / source_id
        existing = self._read_metadata(source_dir)
        if existing:
            return self.get_source(source_id)
        source_dir.mkdir(parents=True, exist_ok=True)

        format_name, raw_name = self._format_from_response(final_url, content_type, response.content)
        raw_path = source_dir / raw_name
        raw_path.write_bytes(response.content)

        title = ""
        if format_name == "html":
            soup = BeautifulSoup(response.text, "html.parser")
            title = first_non_empty(soup.title.string if soup.title else "", Path(urlparse(final_url).path).stem.replace("-", " "))
        else:
            title = Path(urlparse(final_url).path).stem.replace("-", " ")
        title = first_non_empty(title, final_url)

        metadata = {
            "id": source_id,
            "kind": "url",
            "format": format_name,
            "title": title,
            "original_name": raw_name,
            "source_url": source_url,
            "canonical_url": final_url,
            "imported_at": iso_now(),
            "processed_at": "",
            "status": "imported",
            "sha256": sha256_bytes(response.content),
            "mime_type": content_type or "application/octet-stream",
            "raw_items": [
                {
                    "label": "Fetched artifact",
                    "path": str(raw_path.relative_to(self.workspace_root)),
                    "size": raw_path.stat().st_size,
                },
                {
                    "label": "Source URL",
                    "path": final_url,
                    "size": 0,
                },
            ],
            "processing_error": "",
        }
        self._write_metadata(source_dir, metadata)
        return self.process_source(source_id)

    def process_source(self, source_id: str) -> dict:
        source_dir = self.raw_root / source_id
        metadata = self._read_metadata(source_dir)
        if not metadata:
            raise FileNotFoundError(f"Source not found: {source_id}")

        try:
            processed = self._extract_source(metadata)
            processed_dir = self.processed_root / source_id
            processed_dir.mkdir(parents=True, exist_ok=True)

            document_path = processed_dir / "document.md"
            document_path.write_text(
                self._build_processed_document(metadata, processed),
                encoding="utf-8",
            )
            meta_path = processed_dir / "meta.json"
            meta_path.write_text(json.dumps({
                "id": source_id,
                "title": processed["title"],
                "word_count": processed["word_count"],
                "format": metadata["format"],
                "processed_at": iso_now(),
            }, indent=2), encoding="utf-8")

            metadata["title"] = processed["title"]
            metadata["processed_at"] = iso_now()
            metadata["status"] = "processed"
            metadata["processing_error"] = ""
            self._write_metadata(source_dir, metadata)
            return self.get_source(source_id)
        except Exception as exc:
            metadata["status"] = "error"
            metadata["processing_error"] = str(exc)
            self._write_metadata(source_dir, metadata)
            raise

    def _read_metadata(self, source_dir: Path) -> dict:
        path = source_dir / "source.json"
        if not path.exists():
            return {}
        return json.loads(path.read_text(encoding="utf-8"))

    def _write_metadata(self, source_dir: Path, metadata: dict) -> None:
        path = source_dir / "source.json"
        path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")

    def _summarize_metadata(self, metadata: dict) -> dict:
        processed_meta = self._read_processed_meta(metadata["id"])
        processed_path = self._processed_document_path(metadata["id"])
        return {
            "id": metadata["id"],
            "title": metadata["title"],
            "kind": metadata["kind"],
            "format": metadata["format"],
            "status": metadata.get("status", "imported"),
            "original_name": metadata.get("original_name", ""),
            "source_url": metadata.get("canonical_url") or metadata.get("source_url") or "",
            "imported_at": metadata.get("imported_at", ""),
            "processed_at": metadata.get("processed_at", ""),
            "word_count": processed_meta.get("word_count", 0),
            "raw_items": metadata.get("raw_items", []),
            "processed_path": str(processed_path.relative_to(self.workspace_root)) if processed_path.exists() else "",
            "processing_error": metadata.get("processing_error", ""),
        }

    def _processed_document_path(self, source_id: str) -> Path:
        return self.processed_root / source_id / "document.md"

    def _read_processed_meta(self, source_id: str) -> dict:
        path = self.processed_root / source_id / "meta.json"
        if not path.exists():
            return {}
        return json.loads(path.read_text(encoding="utf-8"))

    def _source_primary_file(self, metadata: dict) -> Path:
        raw_items = metadata.get("raw_items", [])
        for item in raw_items:
            path = Path(item["path"])
            if path.is_absolute():
                return path
            candidate = self.workspace_root / path
            if candidate.exists():
                return candidate
        raise FileNotFoundError(f"No raw artifact found for {metadata['id']}")

    def _extract_source(self, metadata: dict) -> dict:
        source_format = metadata["format"]
        path = self._source_primary_file(metadata)
        if source_format == "pdf":
            return self._extract_pdf(metadata, path)
        if source_format == "epub":
            return self._extract_epub(metadata, path)
        if source_format == "docx":
            return self._extract_docx(metadata, path)
        if source_format in {"text", "markdown"}:
            return self._extract_textlike(metadata, path, preserve_markdown=(source_format == "markdown"))
        if source_format == "html":
            return self._extract_html(metadata, path)
        if source_format == "json":
            return self._extract_json(metadata, path)
        if source_format in {"csv", "tsv"}:
            return self._extract_delimited(metadata, path, delimiter="," if source_format == "csv" else "\t")
        raise ValueError(f"Unsupported processed format: {source_format}")

    def _extract_pdf(self, metadata: dict, path: Path) -> dict:
        reader = PdfReader(str(path))
        sections = []
        for index, page in enumerate(reader.pages, start=1):
            text = ensure_text(page.extract_text() or "")
            if not text:
                continue
            sections.append(f"## Page {index}\n\n{text}")
        body = "\n\n".join(sections).strip() or "_No extractable text found in the PDF._"
        return self._result(metadata, body, raw_preview=body)

    def _extract_epub(self, metadata: dict, path: Path) -> dict:
        book = epub.read_epub(str(path))
        sections = []
        preview_parts = []
        for item in book.get_items():
            if item.get_type() != ITEM_DOCUMENT:
                continue
            html = item.get_content().decode("utf-8", errors="ignore")
            section_title = ""
            soup = BeautifulSoup(html, "html.parser")
            heading = soup.find(["h1", "h2", "title"])
            if heading:
                section_title = ensure_text(heading.get_text(" ", strip=True))
            section_title = first_non_empty(section_title, Path(item.get_name()).stem.replace("_", " "))
            markdown = clean_markdown(html_to_markdown(html, heading_style="ATX"))
            if not markdown:
                continue
            sections.append(f"## {section_title}\n\n{markdown}")
            preview_parts.append(html_to_text_preview(html))
        body = "\n\n".join(sections).strip() or "_No extractable text found in the EPUB._"
        preview = "\n\n".join(part for part in preview_parts if part).strip()
        return self._result(metadata, body, raw_preview=preview or body)

    def _extract_docx(self, metadata: dict, path: Path) -> dict:
        document = Document(str(path))
        chunks = []
        for paragraph in document.paragraphs:
            text = ensure_text(paragraph.text)
            if text:
                chunks.append(text)
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([ensure_text(cell.text) for cell in row.cells])
            if rows:
                chunks.append(markdown_table(rows, max_rows=12))
        body = "\n\n".join(chunks).strip() or "_No extractable text found in the DOCX._"
        return self._result(metadata, body, raw_preview=body)

    def _extract_textlike(self, metadata: dict, path: Path, preserve_markdown: bool) -> dict:
        text = decode_text_file(path)
        if preserve_markdown:
            body = text or "_The markdown source is empty._"
        else:
            body = text or "_The text source is empty._"
        return self._result(metadata, body, raw_preview=text)

    def _extract_html(self, metadata: dict, path: Path) -> dict:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        title = first_non_empty(
            soup.title.get_text(" ", strip=True) if soup.title else "",
            metadata.get("title", ""),
        )
        markdown = clean_markdown(html_to_markdown(html, heading_style="ATX"))
        preview = html_to_text_preview(html)
        return self._result(metadata, markdown or "_No visible text found in the HTML source._", raw_preview=preview, title=title)

    def _extract_json(self, metadata: dict, path: Path) -> dict:
        raw_text = decode_text_file(path)
        try:
            payload = json.loads(raw_text)
            pretty = json.dumps(payload, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            pretty = raw_text
        body = "```json\n" + pretty.strip() + "\n```"
        return self._result(metadata, body, raw_preview=pretty)

    def _extract_delimited(self, metadata: dict, path: Path, delimiter: str) -> dict:
        raw_text = decode_text_file(path)
        rows = list(csv.reader(raw_text.splitlines(), delimiter=delimiter))
        sample = markdown_table(rows, max_rows=12)
        lang = "csv" if delimiter == "," else "tsv"
        body = "\n\n".join([
            "## Sample",
            sample,
            "## Full Content",
            f"```{lang}\n{raw_text.strip()}\n```",
        ])
        return self._result(metadata, body, raw_preview=raw_text)

    def _result(self, metadata: dict, body: str, raw_preview: str, title: str | None = None) -> dict:
        final_title = first_non_empty(title or "", metadata.get("title", ""), metadata.get("original_name", "Source"))
        preview = ensure_text(raw_preview)[:TEXT_PREVIEW_LIMIT]
        word_count = len(re.findall(r"\b\w+\b", body))
        return {
            "title": final_title,
            "body": body.strip(),
            "raw_preview": preview,
            "word_count": word_count,
        }

    def _build_processed_document(self, metadata: dict, processed: dict) -> str:
        lines = [
            f"# {processed['title']}",
            "",
            "## Source Record",
            f"- Source ID: `{metadata['id']}`",
            f"- Kind: {metadata['kind']}",
            f"- Format: {metadata['format']}",
            f"- Imported: {metadata.get('imported_at', '')}",
        ]
        if metadata.get("canonical_url"):
            lines.append(f"- Canonical URL: {metadata['canonical_url']}")
        if metadata.get("original_name"):
            lines.append(f"- Original Name: {metadata['original_name']}")
        raw_items = metadata.get("raw_items", [])
        if raw_items:
            lines.extend([
                "",
                "## Raw Artifacts",
                *[
                    f"- {item['label']}: `{item['path']}`"
                    for item in raw_items
                ],
            ])
        lines.extend([
            "",
            "## Canonical Content",
            "",
            processed["body"],
            "",
        ])
        return "\n".join(lines).strip() + "\n"

    def _raw_preview(self, metadata: dict) -> str:
        try:
            return self._extract_source(metadata)["raw_preview"]
        except Exception:
            error = metadata.get("processing_error", "")
            return error or "Raw preview is not available for this source."

    def _parse_windows_url_shortcut(self, path: Path) -> str:
        text = decode_text_file(path)
        for line in text.splitlines():
            if line.upper().startswith("URL="):
                return line.split("=", 1)[1].strip()
        raise ValueError(f"Could not find URL= entry in {path.name}")

    def _format_from_response(self, url: str, content_type: str, payload: bytes) -> tuple[str, str]:
        parsed = urlparse(url)
        suffix = Path(parsed.path).suffix.lower()
        if suffix in SUPPORTED_EXTENSIONS and suffix != ".url":
            format_name = guess_format_from_path(Path(f"source{suffix}"))
            return format_name, f"snapshot{suffix}"
        if any(content_type.startswith(prefix) for prefix in HTML_MIME_PREFIXES):
            return "html", "snapshot.html"
        if any(content_type.startswith(prefix) for prefix in JSON_MIME_PREFIXES):
            return "json", "snapshot.json"
        if any(content_type.startswith(prefix) for prefix in TEXT_MIME_PREFIXES):
            if "csv" in content_type:
                return "csv", "snapshot.csv"
            if "tab-separated-values" in content_type:
                return "tsv", "snapshot.tsv"
            return "text", "snapshot.txt"
        if any(content_type.startswith(prefix) for prefix in PDF_MIME_PREFIXES):
            return "pdf", "snapshot.pdf"

        sniff = payload[:256].decode("utf-8", errors="ignore").lower()
        if "<html" in sniff:
            return "html", "snapshot.html"
        return "text", "snapshot.txt"
